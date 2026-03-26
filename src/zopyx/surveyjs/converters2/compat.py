"""Compatibility layer: adapts converters2 to old converters API."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from zopyx.surveyjs.converters2.types import (
    Attachment,
    Cell,
    CellAddress,
    CellType,
    Response,
    ValueType,
)
from zopyx.surveyjs.converters2.builder import ResponseBuilder
from zopyx.surveyjs.converters2.common import format_datetime


class MockAttachment:
    """Mimics Attachment from old API."""

    def __init__(self, attachment: Attachment):
        self._attachment = attachment

    @property
    def name(self) -> str:
        return self._attachment.name

    @property
    def content(self) -> bytes:
        return self._attachment.content

    @property
    def content_type(self) -> Optional[str]:
        return self._attachment.content_type


class MockItem:
    """Mimics Item from old API."""

    def __init__(self, question_key: str, cells: List[Cell]):
        self.key = question_key
        self._cells = cells

        # Get label from first cell
        self.label = cells[0].label if cells else question_key

    @property
    def value(self) -> Any:
        """Return primary value."""
        if len(self._cells) == 1:
            return self._cells[0].value
        return [c.value for c in self._cells]

    @property
    def values(self) -> List[Any]:
        """Return all values."""
        return [c.value for c in self._cells]

    @property
    def table(self) -> Optional[List[List[Any]]]:
        """Return table data if applicable."""
        # Check if any cells have row_index
        dynamic_cells = [c for c in self._cells if c.address.row_index is not None]
        if not dynamic_cells:
            return None

        # Group by row
        rows = {}
        for cell in dynamic_cells:
            idx = cell.address.row_index or 0
            if idx not in rows:
                rows[idx] = {}
            rows[idx][cell.address.sub_key] = str(cell.value)

        # Get headers
        headers = sorted(
            set(k for row in rows.values() for k in row.keys() if k is not None)
        )

        # Build table
        table = [["#"] + headers]
        for idx in sorted(rows.keys()):
            row = rows[idx]
            table.append([str(idx + 1)] + [row.get(h, "") for h in headers])

        return table

    @property
    def attachments(self) -> List[MockAttachment]:
        """Return attachments from file cells."""
        # Note: Actual attachments are stored in Response.attachments
        # This property is a placeholder for API compatibility
        return []


class SurveyConverter:
    """Mimics SurveyConverter from old API."""

    def __init__(
        self,
        data_path: Path,
        form_path: Path,
        output_dir: Path,
    ):
        """
        Initialize converter.

        Args:
            data_path: Path to JSON data file
            form_path: Path to form schema JSON file
            output_dir: Output directory for exports
        """
        self.data_path = data_path
        self.form_path = form_path
        self.output_dir = output_dir
        self._schema = self._load_schema()

    def _load_schema(self) -> Dict[str, Any]:
        """Load form schema from file."""
        return json.loads(self.form_path.read_text(encoding="utf-8"))

    def load_first_entry(self) -> Dict[str, Any]:
        """Load survey data and return the first entry."""
        payload = json.loads(self.data_path.read_text(encoding="utf-8"))
        if isinstance(payload, list):
            if not payload:
                raise ValueError("Result payload is empty.")
            return payload[0]
        if isinstance(payload, dict):
            return payload
        raise TypeError("Unexpected JSON structure.")

    def collect_items(
        self, entry: Dict[str, Any], poll_id: str
    ) -> Tuple[List[MockItem], List[MockAttachment]]:
        """Collect items and attachments from survey result.
        
        Args:
            entry: Survey result data (the 'result' dict)
            poll_id: Poll/survey ID for context
            
        Returns:
            Tuple of (items, attachments)
        """
        builder = ResponseBuilder(self._schema or {})
        response = builder.build_from_json(entry, poll_id, None, None)

        # Group cells by question key
        by_question: Dict[str, List[Cell]] = {}
        for cell in response.cells:
            key = cell.address.question_key
            if key not in by_question:
                by_question[key] = []
            by_question[key].append(cell)

        items = [MockItem(k, cells) for k, cells in by_question.items()]
        attachments = [MockAttachment(att) for att in response.attachments]
        
        return items, attachments

    def save_attachments(
        self,
        attachments: List[MockAttachment],
    ) -> List[Path]:
        """Save attachments to output directory."""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        saved = []
        for att in attachments:
            path = self.output_dir / att.name
            with open(path, "wb") as f:
                f.write(att.content)
            saved.append(path)

        return saved

    def inline_html_images(
        self, html_body: str, attachments: List[MockAttachment]
    ) -> str:
        """Swap local image references for data URLs so HTML/PDF embed images."""
        from zopyx.surveyjs.converters.image_utils import inline_html_images as _inline
        
        # Convert MockAttachment back to Attachment-like objects
        class AttWrapper:
            def __init__(self, att):
                self.name = att.name
                self.content = att.content
                self.content_type = att.content_type
        
        return _inline(html_body, [AttWrapper(att) for att in attachments])

    def send_email_mailhost(
        self,
        mto: str,
        mfrom: str,
        subject: str,
        message: str,
        attachments: Optional[List[Tuple[str, bytes]]] = None,
    ) -> bool:
        """Send email via Plone MailHost (requires context)."""
        # This would need proper Plone integration
        raise NotImplementedError("send_email_mailhost requires Plone context")


def build_markdown(
    items: List[MockItem],
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> str:
    """Build markdown export (compatibility wrapper)."""
    from zopyx.surveyjs.converters2.markdown import build_markdown as _build

    # Convert back to Response format for internal functions
    response = _items_to_response(items, poll_id, creator, created)
    return _build(response)


def build_html(markdown_body: str, attachments: List[MockAttachment]) -> str:
    """Build HTML from markdown (compatibility wrapper)."""
    from markdown2 import markdown

    html_content = markdown(
        markdown_body,
        extras=[
            "tables",
            "fenced-code-blocks",
            "header-ids",
            "strike",
            "task_list",
        ],
    )

    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Survey Export</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 900px;
            margin: 2em auto;
            padding: 0 1em;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            border-bottom: 2px solid #4a90d9;
            padding-bottom: 0.3em;
        }}
        h2 {{
            margin-top: 1.5em;
            color: #2c5282;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #e2e8f0;
            padding: 0.5em;
            text-align: left;
        }}
        th {{
            background: #f7fafc;
            font-weight: 600;
        }}
        tr:nth-child(even) {{
            background: #f7fafc;
        }}
        code {{
            background: #edf2f7;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-size: 0.9em;
        }}
        pre {{
            background: #f7fafc;
            padding: 1em;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #cbd5e0;
            margin: 0;
            padding-left: 1em;
            color: #4a5568;
        }}
        .metadata {{
            color: #666;
            font-size: 0.9em;
            margin-bottom: 2em;
        }}
        .attachment {{
            color: #666;
            font-style: italic;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""


def write_pdf(
    html_body: str,
    destination: Path,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Path:
    """Write PDF export (compatibility wrapper)."""
    from weasyprint import HTML

    destination.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html_body).write_pdf(destination)
    return destination


def build_table_rows(items: List[MockItem]) -> List[Dict[str, Any]]:
    """Build table rows from items for CSV/XLSX export."""
    rows = []
    for item in items:
        if item.table:
            # For tables, add label and flattened rows
            for i, row in enumerate(item.table):
                if i == 0:  # Header row
                    continue
                row_dict = {"_question": item.label, "_row": i}
                for j, val in enumerate(row):
                    header = item.table[0][j] if j < len(item.table[0]) else f"col_{j}"
                    row_dict[header] = val
                rows.append(row_dict)
        else:
            # Simple values
            for val in item.values:
                rows.append({"_question": item.label, "_key": item.key, "value": val})
    return rows


def write_csv(
    table_rows: List[Dict[str, Any]],
    destination: Path,
) -> Path:
    """Write CSV export from pre-built table rows."""
    import csv

    destination.parent.mkdir(parents=True, exist_ok=True)
    
    if not table_rows:
        with open(destination, "w", newline="", encoding="utf-8") as f:
            f.write("")
        return destination
    
    # Get all fieldnames
    fieldnames = set()
    for row in table_rows:
        fieldnames.update(row.keys())
    fieldnames = sorted(fieldnames)
    
    with open(destination, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(table_rows)
    
    return destination


def write_xlsx(
    table_rows: List[Dict[str, Any]],
    destination: Path,
) -> Path:
    """Write XLSX export from pre-built table rows."""
    destination.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        from openpyxl import Workbook
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Survey Results"
        
        if not table_rows:
            wb.save(destination)
            return destination
        
        # Get headers
        fieldnames = set()
        for row in table_rows:
            fieldnames.update(row.keys())
        fieldnames = sorted(fieldnames)
        
        # Write header
        ws.append(fieldnames)
        
        # Write data
        for row in table_rows:
            ws.append([row.get(k, "") for k in fieldnames])
        
        wb.save(destination)
    except ImportError:
        # Fallback: create empty file
        destination.write_text("", encoding="utf-8")
    
    return destination


def write_html(
    markdown_body: str,
    attachments: List[MockAttachment],
    destination: Path,
) -> Path:
    """Write HTML export from markdown body."""
    from markdown2 import markdown

    html_content = markdown(
        markdown_body,
        extras=[
            "tables",
            "fenced-code-blocks",
            "header-ids",
            "strike",
            "task_list",
        ],
    )

    # Build full HTML page
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Survey Export</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 900px;
            margin: 2em auto;
            padding: 0 1em;
            line-height: 1.6;
            color: #333;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #e2e8f0;
            padding: 0.5em;
            text-align: left;
        }}
        th {{
            background: #f7fafc;
            font-weight: 600;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(full_html, encoding="utf-8")
    return destination


def write_docx(
    items: List[MockItem],
    destination: Path,
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Path:
    """Write compact DOCX export (compatibility wrapper)."""
    from docx import Document

    doc = Document()

    # Compact header - smaller heading
    doc.add_heading(f"Survey: {poll_id}", level=2)

    # Metadata on single lines
    if creator:
        p = doc.add_paragraph()
        p.add_run("By: ").bold = True
        p.add_run(creator)
    if created:
        p = doc.add_paragraph()
        p.add_run("On: ").bold = True
        p.add_run(format_datetime(created))

    # Items in compact format
    for item in items:
        if item.table:
            # Table data - label on separate line as bold
            p = doc.add_paragraph()
            p.add_run(item.label).bold = True

            # Create table
            table = doc.add_table(rows=len(item.table), cols=len(item.table[0]))
            table.style = "Table Grid"
            for row_idx, row in enumerate(item.table):
                for col_idx, cell in enumerate(row):
                    table.rows[row_idx].cells[col_idx].text = str(cell)

            # Bold header row
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        else:
            # Simple fields - label and value on same line
            for val in item.values:
                p = doc.add_paragraph()
                p.add_run(f"{item.label}").bold = True
                p.add_run(f": {val}")

        # Compact attachment list
        if item.attachments:
            for att in item.attachments:
                p = doc.add_paragraph()
                p.add_run(f"📎 {att.name}").italic = True

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)
    return destination


def write_xml(
    items: List[MockItem],
    destination: Path,
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Path:
    """Write XML export (compatibility wrapper)."""
    from zopyx.surveyjs.converters2.xml_export import write_xml as _write

    response = _items_to_response(items, poll_id, creator, created)
    return _write(response, destination)


def write_json(
    items: List[MockItem],
    destination: Path,
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Path:
    """Write JSON export (compatibility wrapper)."""
    import json

    destination.parent.mkdir(parents=True, exist_ok=True)

    data = {
        "poll_id": poll_id,
        "creator": creator,
        "created": created,
        "items": [
            {
                "key": item.key,
                "label": item.label,
                "value": item.value,
                "values": item.values,
                "table": item.table,
            }
            for item in items
        ],
    }

    with open(destination, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False, default=str)

    return destination


def write_text(
    items: List[MockItem],
    destination: Path,
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Path:
    """Write plain text export (compatibility wrapper)."""
    destination.parent.mkdir(parents=True, exist_ok=True)

    lines = [f"Survey Response: {poll_id}", "=" * 50, ""]

    if creator:
        lines.append(f"Created by: {creator}")
    if created:
        lines.append(f"Created on: {format_datetime(created)}")
    if creator or created:
        lines.append("")

    for item in items:
        lines.append(f"{item.label} ({item.key})")
        lines.append("-" * 40)

        if item.table:
            for row in item.table:
                lines.append(" | ".join(str(c) for c in row))
        else:
            for val in item.values:
                lines.append(f"  - {val}")

        for att in item.attachments:
            lines.append(f"  [Attachment: {att.name}]")

        lines.append("")

    with open(destination, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return destination


def write_markdown(
    items: List[MockItem],
    destination: Path,
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Path:
    """Write markdown export directly to file."""
    markdown_body = build_markdown(items, poll_id, creator, created)
    destination.parent.mkdir(parents=True, exist_ok=True)
    with open(destination, "w", encoding="utf-8") as f:
        f.write(markdown_body)
    return destination


def _items_to_response(
    items: List[MockItem],
    poll_id: str,
    creator: Optional[str] = None,
    created: Optional[str] = None,
) -> Response:
    """Convert MockItems back to Response format."""
    # This is used to bridge between old and new APIs
    cells = []
    attachments = []

    for item in items:
        for i, val in enumerate(item.values):
            addr = CellAddress(question_key=item.key)
            if item.table and i > 0:  # Skip header row
                addr = CellAddress(question_key=item.key, row_index=i - 1)

            cell_type = CellType.SCALAR
            if item.attachments:
                for att in item.attachments:
                    cell_type = CellType.FILE
                    attachments.append(
                        Attachment(
                            attachment_id=f"{poll_id}_{att.name}",
                            name=att.name,
                            content_type=att.content_type,
                            content=att.content,
                        )
                    )

            cells.append(
                Cell(
                    address=addr,
                    value=val,
                    field_type="unknown",
                    cell_type=cell_type,
                    value_type=ValueType.STRING,
                    label=item.label,
                    column_name="",
                )
            )

    return Response(
        response_id=poll_id,
        creator=creator,
        created=created,
        cells=cells,
        attachments=attachments,
    )
