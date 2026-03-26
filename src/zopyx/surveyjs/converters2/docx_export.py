"""DOCX converter for Response objects - compact format."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .common import format_datetime
from .types import CellType, Response


def write_docx(response: Response, destination: Path) -> Path:
    """Write a compact DOCX export."""
    destination.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    # Compact header
    doc.add_heading(f"Survey: {response.response_id}", level=2)
    
    # Metadata on single lines
    if response.creator:
        p = doc.add_paragraph()
        p.add_run("By: ").bold = True
        p.add_run(response.creator)
    
    if response.created:
        p = doc.add_paragraph()
        p.add_run("On: ").bold = True
        p.add_run(format_datetime(response.created))
    
    # Group cells by question
    by_question = {}
    for cell in response.cells:
        key = cell.address.question_key
        if key not in by_question:
            by_question[key] = []
        by_question[key].append(cell)
    
    # Render each question
    for question_key, cells in by_question.items():
        schema = response.question_schemas.get(question_key)
        title = schema.title if schema else question_key
        
        # Check if has dynamic content (table)
        has_dynamic = any(c.address.row_index is not None for c in cells)
        
        if has_dynamic:
            # Question label as bold text
            p = doc.add_paragraph()
            p.add_run(title).bold = True
            
            # Build table
            by_row = {}
            for cell in cells:
                idx = cell.address.row_index or 0
                if idx not in by_row:
                    by_row[idx] = {}
                by_row[idx][cell.address.sub_key] = str(cell.value)
            
            # Get headers from schema or data
            if schema and schema.columns:
                headers = [c.get("title", c.get("name")) for c in schema.columns]
            else:
                all_keys = set()
                for row_data in by_row.values():
                    all_keys.update(row_data.keys())
                headers = sorted(all_keys)
            
            # Create compact table
            table = doc.add_table(rows=len(by_row) + 1, cols=len(headers))
            table.style = "Table Grid"
            
            # Header row
            for col_idx, header in enumerate(headers):
                table.rows[0].cells[col_idx].text = header
            
            # Data rows
            for row_idx, idx in enumerate(sorted(by_row.keys()), 1):
                row_data = by_row[idx]
                for col_idx, header in enumerate(headers):
                    table.rows[row_idx].cells[col_idx].text = row_data.get(header, "")
            
            # Bold header
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        else:
            # Simple fields - label and value on same line
            for cell in cells:
                p = doc.add_paragraph()
                if cell.address.sub_key:
                    # Matrix/multipletext item - label in brackets
                    p.add_run(f"{cell.label}").bold = True
                    p.add_run(f": {cell.value}")
                else:
                    # Simple field
                    p.add_run(f"{title}").bold = True
                    if cell.cell_type == CellType.FILE:
                        p.add_run(f": {cell.display_value or str(cell.value)}")
                    else:
                        p.add_run(f": {cell.value}")
    
    # Compact attachment list
    if response.attachments:
        p = doc.add_paragraph()
        p.add_run("Attachments: ").bold = True
        att_names = ", ".join(att.name for att in response.attachments)
        p.add_run(att_names)
    
    doc.save(destination)
    return destination
