"""DOCX converter."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document

from .types import Item


def write_docx(
    items: Iterable[Item],
    destination: Path,
    poll_id: str,
    creator: str | None = None,
    created: str | None = None,
) -> Path:
    doc = Document()

    doc.add_heading(f"Survey Response: {poll_id}", level=1)

    if creator:
        p = doc.add_paragraph()
        p.add_run("Created by: ").bold = True
        p.add_run(creator)

    if created:
        from datetime import datetime
        try:
            dt = datetime.fromisoformat(created.replace("Z", "+00:00"))
            formatted_date = dt.strftime("%B %d, %Y at %I:%M %p %Z")
        except (ValueError, AttributeError):
            formatted_date = created
        p = doc.add_paragraph()
        p.add_run("Created on: ").bold = True
        p.add_run(formatted_date)

    if creator or created:
        doc.add_paragraph()

    for item in items:
        doc.add_heading(f"{item.label} ({item.key})", level=2)
        if item.table:
            table = doc.add_table(rows=len(item.table), cols=len(item.table[0]))
            table.style = "Light Grid Accent 1"
            for row_idx, row in enumerate(item.table):
                row_cells = table.rows[row_idx].cells
                for col_idx, cell in enumerate(row):
                    row_cells[col_idx].text = cell
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        else:
            for val in item.values:
                doc.add_paragraph(f"- {val}")
        for att in item.attachments:
            desc = att.content_type or "binary"
            doc.add_paragraph(f"Attachment: {att.name} ({desc})")
        doc.add_paragraph()

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)
    return destination
