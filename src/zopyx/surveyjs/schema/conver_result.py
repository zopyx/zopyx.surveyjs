#!/usr/bin/env -S uv run
# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "markdown2>=2.4.12",
#     "weasyprint>=62.3",
#     "openpyxl>=3.1.3",
#     "python-docx>=1.1.0",
# ]
# ///
from __future__ import annotations

"""SurveyJS result converter producing multiple output formats with attachment handling."""

import argparse
import base64
import csv
import email.utils
import json
import logging
import mimetypes
import os
import smtplib
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from email.message import EmailMessage
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence, Tuple

from docx import Document
from docx.shared import Inches, Pt
from markdown2 import markdown
from openpyxl import Workbook
from weasyprint import HTML

ROOT = Path(__file__).parent
SURVEY_DATA_PATH = ROOT / "survey-data-form.json"
FORM_PATH = ROOT / "survey-form-form.json"
OUTPUT_DIR = ROOT / "output"

logger = logging.getLogger(__name__)

# Environment variable keys used for configuration overrides.
ENV_DOTENV_PATH = "SURVEY_DOTENV_PATH"
ENV_DATA_PATH = "SURVEYJS_DATA_JSON"
ENV_FORM_PATH = "SURVEYJS_FORM_JSON"
ENV_EMAIL_RECIPIENT = "SURVEY_EMAIL_RECIPIENT"
ENV_SMTP_HOST = "SURVEY_SMTP_HOST"
ENV_SMTP_PORT = "SURVEY_SMTP_PORT"
ENV_SMTP_USERNAME = "SURVEY_SMTP_USERNAME"
ENV_SMTP_PASSWORD = "SURVEY_SMTP_PASSWORD"
ENV_SMTP_STARTTLS = "SURVEY_SMTP_STARTTLS"
ENV_EMAIL_SENDER = "SURVEY_EMAIL_SENDER"


def data_default() -> str:
    """Resolve survey data path using env override when provided."""
    return os.environ.get(ENV_DATA_PATH, str(SURVEY_DATA_PATH))


def form_default() -> str:
    """Resolve survey form path using env override when provided."""
    return os.environ.get(ENV_FORM_PATH, str(FORM_PATH))


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    """Parse CLI options for selecting input and output settings."""
    load_dotenv()

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--data",
        default=data_default(),
        help=(
            f"Path to survey data JSON "
            f"(default: {SURVEY_DATA_PATH}, env override: ${ENV_DATA_PATH})."
        ),
    )
    parser.add_argument(
        "--form",
        default=form_default(),
        help=(
            f"Path to survey form JSON "
            f"(default: {FORM_PATH}, env override: ${ENV_FORM_PATH})."
        ),
    )
    parser.add_argument(
        "--output",
        default=str(OUTPUT_DIR),
        help=f"Output directory (default: {OUTPUT_DIR}).",
    )
    parser.add_argument(
        "--formats",
        default="all",
        help="Comma-separated formats to emit (text,md,html,pdf,csv,xlsx,xml,docx,json). Default: all.",
    )
    parser.add_argument(
        "--email",
        default=os.environ.get(ENV_EMAIL_RECIPIENT),
        help=(
            "Email recipient to receive all generated files as attachments "
            f"(default from ${ENV_EMAIL_RECIPIENT})."
        ),
    )
    return parser.parse_args()


def parse_formats(spec: str) -> set[str]:
    """Normalize and validate requested formats."""
    allowed = {"text", "md", "html", "pdf", "csv", "xlsx", "xml", "docx", "json"}
    if spec.lower() == "all":
        return allowed
    requested = {part.strip().lower() for part in spec.split(",") if part.strip()}
    invalid = requested - allowed
    if invalid:
        raise ValueError(f"Unknown formats: {', '.join(sorted(invalid))}")
    return requested or allowed


def slugify(value: Any) -> str:
    """Return a filesystem-safe slug for IDs or filenames."""
    text = str(value)
    return "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in text).strip("_") or "sample"


def load_dotenv() -> None:
    """Populate environment variables from a .env file if present."""
    dotenv_path = Path(os.environ.get(ENV_DOTENV_PATH, ROOT / ".env"))
    if not dotenv_path.exists():
        logger.info("No .env file found at %s", dotenv_path)
        return

    logger.info("Loading .env file from %s", dotenv_path)
    loaded_vars = []
    for line in dotenv_path.read_text().splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value
            loaded_vars.append(key)

    if loaded_vars:
        logger.info("Loaded environment variables: %s", ", ".join(loaded_vars))


@dataclass
class Attachment:
    name: str
    content: bytes
    content_type: str | None = None
    field_label: str | None = None

    @property
    def is_image(self) -> bool:
        return bool(self.content_type and self.content_type.startswith("image/"))

    def data_url(self) -> str:
        if not self.content_type:
            mime = mimetypes.guess_type(self.name)[0]
            ctype = mime or "application/octet-stream"
        else:
            ctype = self.content_type
        encoded = base64.b64encode(self.content).decode("ascii")
        return f"data:{ctype};base64,{encoded}"


@dataclass
class Item:
    key: str
    label: str
    values: List[str]
    attachments: List[Attachment]
    field_type: str | None = None
    raw_value: Any | None = None
    table: List[List[str]] | None = None
    table_columns: List[Tuple[str, str]] | None = None


class SurveyConverter:
    """Core converter for SurveyJS data to multiple export formats."""

    def __init__(self, data_path: Path, form_path: Path, output_dir: Path) -> None:
        self.data_path = data_path
        self.form_path = form_path
        self.output_dir = output_dir
        self.schema = self.load_schema()

    def load_first_entry(self) -> Dict[str, Any]:
        """Load survey data and return the first entry (dict) from a list or a single dict."""
        payload = json.loads(self.data_path.read_text())
        if isinstance(payload, list):
            if not payload:
                raise ValueError("Result payload is empty.")
            if not isinstance(payload[0], dict):
                raise TypeError("First entry is not a JSON object.")
            return payload[0]
        if isinstance(payload, dict):
            return payload
        raise TypeError("Unexpected JSON structure.")

    def load_schema(self) -> Dict[str, Dict[str, Any]]:
        """Index form elements by name for quick label lookup."""
        schema = json.loads(self.form_path.read_text())
        elements: Dict[str, Dict[str, Any]] = {}
        for page in schema.get("pages", []):
            for element in page.get("elements", []):
                name = element.get("name")
                if not name:
                    continue
                elements[name] = element
        return elements

    def extract_poll_id(self, entry: Dict[str, Any]) -> str:
        """Pick a poll identifier from common keys, falling back to 'sample'."""
        for key in ("poll_id", "pollId", "id"):
            if key in entry:
                return slugify(entry[key])
        return "sample"

    def base64_from_data_url(self, value: str) -> Tuple[str | None, str]:
        """Split a data URL into content-type and base64 payload; otherwise return the raw value."""
        if not value.startswith("data:") or "," not in value:
            return None, value
        header, encoded = value.split(",", 1)
        meta = header[len("data:") :]
        content_type = meta.split(";")[0] if ";" in meta else meta
        return content_type or None, encoded

    def decode_base64_payload(self, encoded: str) -> bytes | None:
        """Decode base64 content if it looks valid; return None otherwise."""
        payload = encoded.strip()
        if len(payload) < 16 or len(payload) % 4 != 0:
            return None
        try:
            return base64.b64decode(payload, validate=True)
        except Exception:
            return None

    def extract_attachments(self, name: str, label: str, value: Any, poll_id: str) -> Tuple[List[str], List[Attachment]]:
        """Collect decoded attachments from file fields and return human-friendly value lines."""
        lines: List[str] = []
        attachments: List[Attachment] = []

        def handle_single(content: str, filename: str | None, content_type: str | None) -> None:
            ctype, encoded = self.base64_from_data_url(content)
            raw = self.decode_base64_payload(encoded)
            if raw is None:
                lines.append(content)
                return
            ext = mimetypes.guess_extension(ctype or content_type or "") or ".bin"
            fname = filename or f"{poll_id}_{name}{ext}"
            attachments.append(Attachment(fname, raw, ctype or content_type, field_label=label))
            lines.append(f"stored attachment: {fname}")

        if isinstance(value, list):
            for idx, item in enumerate(value):
                if isinstance(item, dict):
                    content = item.get("content") or item.get("base64")
                    if not content:
                        continue
                    fname = item.get("name") or f"{poll_id}_{name}_{idx}"
                    handle_single(content, fname, item.get("type"))
                elif isinstance(item, str):
                    handle_single(item, f"{poll_id}_{name}_{idx}", None)
        elif isinstance(value, dict):
            content = value.get("content") or value.get("base64")
            if content:
                handle_single(content, value.get("name"), value.get("type"))
        elif isinstance(value, str):
            handle_single(value, None, None)
        else:
            lines.append(str(value))

        return lines or ["(no file content)"], attachments

    def format_matrix(self, value: Dict[str, Any], element: Dict[str, Any]) -> List[str]:
        """Render matrix answers with row and column labels."""
        rows = {row.get("value"): row.get("text", row.get("value")) for row in element.get("rows", [])}
        cols = {
            str(col.get("value")): col.get("text", col.get("value"))
            for col in element.get("columns", [])
        }
        lines = []
        for row_key, cell_value in value.items():
            row_label = rows.get(row_key, row_key)
            cell_label = cols.get(str(cell_value), cell_value)
            lines.append(f"{row_label}: {cell_label}")
        return lines or [json.dumps(value, ensure_ascii=False)]

    def format_matrixdynamic_table(
        self, value: Any, element: Dict[str, Any]
    ) -> Tuple[List[List[str]], List[Tuple[str, str]], List[Dict[str, Any]]]:
        """Render matrixdynamic answers as a text-friendly table."""
        rows_list: List[Any]
        if isinstance(value, list):
            rows_list = value
        elif isinstance(value, dict):
            rows_list = list(value.values())
        else:
            rows_list = []

        columns = element.get("columns") or []
        col_names: List[str] = []
        col_headers: List[str] = []
        col_meta: List[Tuple[str, str]] = []
        for col in columns:
            name = col.get("name") or col.get("value")
            header = col.get("title") or name
            if name is None:
                continue
            name_str = str(name)
            header_str = str(header or name)
            col_names.append(name_str)
            col_headers.append(header_str)
            col_meta.append((name_str, header_str))

        if not col_names:
            if rows_list and isinstance(rows_list[0], dict):
                keys = []
                for row in rows_list:
                    if isinstance(row, dict):
                        for key in row.keys():
                            if key not in keys:
                                keys.append(key)
                col_names = [str(key) for key in keys] or ["Value"]
                col_headers = col_names[:]
                col_meta = [(name, name) for name in col_names]
            else:
                col_names = ["Value"]
                col_headers = ["Value"]
                col_meta = [("Value", "Value")]

        table: List[List[str]] = [col_headers]
        if not rows_list:
            table.append(["(empty)" for _ in col_names])
            return table, col_meta, []

        for row in rows_list:
            if isinstance(row, dict):
                cells = [self.stringify_cell(row.get(name)) for name in col_names]
            else:
                cells = [self.stringify_cell(row)]
            if len(cells) < len(col_names):
                cells.extend([""] * (len(col_names) - len(cells)))
            table.append(cells)
        raw_rows = [row for row in rows_list if isinstance(row, dict)]
        return table, col_meta, raw_rows

    def stringify_cell(self, value: Any) -> str:
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False)
        return "" if value is None else str(value)

    def format_value(
        self,
        name: str,
        label: str,
        value: Any,
        element: Dict[str, Any],
        poll_id: str,
    ) -> Tuple[List[str], List[Attachment], List[List[str]] | None, List[Tuple[str, str]] | None, Any | None]:
        """Format a single field based on schema type, returning display lines and attachments."""
        if element.get("type") == "file":
            lines, attachments = self.extract_attachments(name, label, value, poll_id)
            return lines, attachments, None, None, None

        if element.get("type") == "matrix" and isinstance(value, dict):
            return self.format_matrix(value, element), [], None, None, None

        if element.get("type") == "matrixdynamic":
            table, table_columns, raw_rows = self.format_matrixdynamic_table(value, element)
            return [json.dumps(value, ensure_ascii=False)], [], table, table_columns, raw_rows

        if isinstance(value, bool):
            return ["Yes" if value else "No"], [], None, None, None

        if isinstance(value, list):
            return ([", ".join(str(v) for v in value)] if value else ["(empty)"]), [], None, None, None

        if isinstance(value, dict):
            return [json.dumps(value, ensure_ascii=False)], [], None, None, None

        return [str(value)], [], None, None, None

    def collect_items(self, entry: Dict[str, Any], poll_id: str) -> Tuple[List[Item], List[Attachment]]:
        """Assemble items with labels, values, and attachments for downstream rendering."""
        items: List[Item] = []
        attachments: List[Attachment] = []
        for name, value in entry.items():
            element = self.schema.get(name, {})
            field_type = element.get("type")
            label = element.get("title") or element.get("name") or name
            lines, extra, table, table_columns, raw_value = self.format_value(
                name, label, value, element, poll_id
            )
            items.append(
                Item(
                    key=name,
                    label=label,
                    field_type=field_type,
                    values=lines,
                    raw_value=raw_value,
                    attachments=extra,
                    table=table,
                    table_columns=table_columns,
                )
            )
            attachments.extend(extra)
        return items, attachments

    def render_text_table(self, table: List[List[str]]) -> List[str]:
        if not table:
            return ["(empty)"]
        col_count = max(len(row) for row in table)
        normalized = [row + [""] * (col_count - len(row)) for row in table]
        widths = [
            max(len(str(row[idx])) for row in normalized)
            for idx in range(col_count)
        ]

        def format_row(row: List[str]) -> str:
            padded = [str(cell).ljust(widths[idx]) for idx, cell in enumerate(row)]
            return " | ".join(padded)

        lines = [format_row(normalized[0])]
        if len(normalized) > 1:
            separator = "-+-".join("-" * width for width in widths)
            lines.append(separator)
            for row in normalized[1:]:
                lines.append(format_row(row))
        return lines

    def render_markdown_table(self, table: List[List[str]]) -> List[str]:
        if not table:
            return ["(empty)"]
        col_count = max(len(row) for row in table)
        normalized = [row + [""] * (col_count - len(row)) for row in table]
        header = "| " + " | ".join(normalized[0]) + " |"
        separator = "| " + " | ".join("---" for _ in range(col_count)) + " |"
        rows = ["| " + " | ".join(row) + " |" for row in normalized[1:]]
        return [header, separator, *rows]

    def build_text(self, items: Iterable[Item]) -> List[str]:
        """Produce a plain-text representation with attachment notices."""
        lines: List[str] = ["Survey response", ""]
        for item in items:
            lines.append(f"{item.label}:")
            if item.table:
                table_lines = self.render_text_table(item.table)
                lines.extend(f"  {line}" for line in table_lines)
            else:
                for val in item.values:
                    lines.append(f"  - {val}")
            for att in item.attachments:
                desc = att.content_type or "binary"
                lines.append(f"  - Attachment: {att.name} ({desc})")
            lines.append("")
        return lines

    def build_markdown(self, items: Iterable[Item], poll_id: str) -> str:
        """Produce Markdown output including image references."""
        parts: List[str] = [f"# Survey response ({poll_id})", ""]
        for item in items:
            parts.append(f"## {item.label} ({item.key})")
            if item.table:
                parts.extend(self.render_markdown_table(item.table))
            else:
                for val in item.values:
                    parts.append(f"- {val}")
            for att in item.attachments:
                if att.is_image:
                    parts.append(f"![{item.label} - {att.name}]({att.name})")
                else:
                    parts.append(f"[{att.name}]({att.name})")
            parts.append("")
        return "\n".join(parts).strip() + "\n"

    def build_table_rows(self, items: Iterable[Item]) -> List[Tuple[str, str, str, str]]:
        """Flatten items into rows for CSV/Excel export."""
        rows: List[Tuple[str, str, str, str]] = []
        for item in items:
            value = "; ".join(item.values)
            att_desc = "; ".join(
                f"{att.name} ({att.content_type or 'binary'})" for att in item.attachments
            )
            rows.append((item.key, item.label, value, att_desc))
        return rows

    def build_xml(self, items: Iterable[Item], poll_id: str) -> str:
        """Generate XML representation of survey results."""
        root = ET.Element("survey_response")
        root.set("poll_id", poll_id)

        for item in items:
            field = ET.SubElement(root, "field")
            field.set("key", item.key)
            field.set("label", item.label)

            if item.table:
                table_elem = ET.SubElement(field, "table")
                columns = item.table_columns or []
                if len(columns) < len(item.table[0]):
                    padding = len(item.table[0]) - len(columns)
                    columns = columns + [("", "")] * padding
                for idx, row in enumerate(item.table):
                    row_elem = ET.SubElement(table_elem, "row")
                    row_elem.set("header", str(idx == 0).lower())
                    for col_idx, cell in enumerate(row):
                        cell_elem = ET.SubElement(row_elem, "cell")
                        if columns:
                            col_key, col_label = columns[col_idx]
                            if col_label:
                                cell_elem.set("label", col_label)
                            if col_key:
                                cell_elem.set("key", col_key.lower())
                        cell_elem.text = cell
            else:
                values_elem = ET.SubElement(field, "values")
                for val in item.values:
                    value_elem = ET.SubElement(values_elem, "value")
                    value_elem.text = val

            # Add attachments if present
            if item.attachments:
                attachments_elem = ET.SubElement(field, "attachments")
                for att in item.attachments:
                    att_elem = ET.SubElement(attachments_elem, "attachment")
                    att_elem.set("name", att.name)
                    if att.content_type:
                        att_elem.set("content_type", att.content_type)
                    att_elem.set("is_image", str(att.is_image).lower())

        # Pretty print the XML with indentation
        ET.indent(root, space="  ")
        return '<?xml version="1.0" encoding="UTF-8"?>\n' + ET.tostring(root, encoding="unicode", method="xml")

    def build_json(
        self,
        items: Iterable[Item],
        poll_id: str,
        creator: str | None = None,
        created: str | None = None,
    ) -> str:
        """Generate JSON representation of survey results."""
        payload = {
            "poll_id": poll_id,
            "creator": creator,
            "created": created,
            "fields": [],
        }
        for item in items:
            values = item.values
            if item.field_type == "matrixdynamic" and item.raw_value is not None:
                values = item.raw_value
            field = {
                "key": item.key,
                "label": item.label,
                "values": item.values,
                "attachments": [
                    {
                        "name": att.name,
                        "content_type": att.content_type,
                        "is_image": att.is_image,
                    }
                    for att in item.attachments
                ],
            }
            field["values"] = values
            if item.table:
                field["table"] = item.table
            if item.table_columns:
                field["table_columns"] = [
                    {"key": key, "label": label} for key, label in item.table_columns
                ]
            payload["fields"].append(field)
        return json.dumps(payload, ensure_ascii=False, indent=2) + "\n"

    def inline_html_images(self, html_body: str, attachments: Iterable[Attachment]) -> str:
        """Swap local image references for data URLs so HTML/PDF embed images."""
        updated = html_body
        for att in attachments:
            if not att.is_image:
                continue
            data_url = att.data_url()
            updated = updated.replace(f'src="{att.name}"', f'src="{data_url}"')
            updated = updated.replace(f"src='{att.name}'", f"src='{data_url}'")
        return updated

    def wrap_pdf_html(self, html_body: str, creator: str = None, created: str = None) -> str:
        """Wrap converted HTML with styles suitable for PDF rendering."""
        style = """
        <style>
          img { max-width: 75%; height: auto; display: block; margin: 0.3em 0; }
          table { border-collapse: collapse; margin: 0.4em 0; }
          th, td { border: 1px solid #ccc; padding: 0.2em 0.4em; }
          body { font-family: sans-serif; font-size: 10pt; line-height: 1.3; }
          h1 { font-size: 14pt; margin: 0.5em 0 0.3em 0; }
          h2 { font-size: 11pt; margin: 0.4em 0 0.2em 0; font-weight: 600; }
          h3 { font-size: 10pt; margin: 0.3em 0 0.2em 0; }
          p { margin: 0.2em 0; }
          ul, ol { margin: 0.2em 0; padding-left: 1.5em; }
          li { margin: 0.1em 0; }
          .metadata { margin-bottom: 0.8em; color: #666; font-size: 9pt; }
          .metadata p { margin: 0.2em 0; }
        </style>
        """

        # Build metadata section
        metadata_html = ""
        if creator or created:
            metadata_parts = []
            if creator:
                metadata_parts.append(f"<p><strong>Created by:</strong> {creator}</p>")
            if created:
                from datetime import datetime
                try:
                    dt = datetime.fromisoformat(created.replace('Z', '+00:00'))
                    formatted_date = dt.strftime("%B %d, %Y at %I:%M %p %Z")
                except (ValueError, AttributeError):
                    formatted_date = created
                metadata_parts.append(f"<p><strong>Created on:</strong> {formatted_date}</p>")
            metadata_html = f'<div class="metadata">{"".join(metadata_parts)}</div>'

        return f"<html><head>{style}</head><body>{metadata_html}{html_body}</body></html>"

    def wrap_html_output(self, html_body: str) -> str:
        """Wrap HTML output with styles for browser rendering."""
        style = """
        <style>
          img { max-width: 1024px; height: auto; display: block; margin: 0.3em 0; }
          table { border-collapse: collapse; margin: 0.4em 0; }
          th, td { border: 1px solid #ccc; padding: 0.2em 0.4em; }
        </style>
        """
        return f"<html><head>{style}</head><body>{html_body}</body></html>"

    def save_pdf(self, html_body: str, destination: Path) -> None:
        """Render HTML into a PDF using WeasyPrint."""
        destination.parent.mkdir(parents=True, exist_ok=True)
        HTML(string=html_body).write_pdf(destination)

    def write_csv(self, rows: List[Tuple[str, str, str, str]], destination: Path) -> None:
        """Write tabular rows to CSV."""
        destination.parent.mkdir(parents=True, exist_ok=True)
        with destination.open("w", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            writer.writerow(["Key", "Field", "Value", "Attachments"])
            writer.writerows(rows)

    def write_excel(self, rows: List[Tuple[str, str, str, str]], destination: Path) -> None:
        """Write tabular rows to an Excel workbook."""
        wb = Workbook()
        ws = wb.active
        ws.title = "Survey"
        ws.append(["Key", "Field", "Value", "Attachments"])
        for row in rows:
            ws.append(list(row))
        destination.parent.mkdir(parents=True, exist_ok=True)
        wb.save(destination)

    def write_docx(self, items: Iterable[Item], destination: Path, poll_id: str, creator: str = None, created: str = None) -> None:
        """Write survey content to a Word document."""
        doc = Document()

        # Add title
        title = doc.add_heading(f"Survey Response: {poll_id}", level=1)

        # Add creator and created information
        if creator:
            p = doc.add_paragraph()
            p.add_run("Created by: ").bold = True
            p.add_run(creator)

        if created:
            from datetime import datetime
            try:
                dt = datetime.fromisoformat(created.replace('Z', '+00:00'))
                formatted_date = dt.strftime("%B %d, %Y at %I:%M %p %Z")
            except (ValueError, AttributeError):
                formatted_date = created
            p = doc.add_paragraph()
            p.add_run("Created on: ").bold = True
            p.add_run(formatted_date)

        # Add spacing before table
        if creator or created:
            doc.add_paragraph()

        for item in items:
            doc.add_heading(f"{item.label} ({item.key})", level=2)
            if item.table:
                table = doc.add_table(rows=len(item.table), cols=len(item.table[0]))
                table.style = "Light Grid Accent 1"
                for row_idx, row in enumerate(item.table):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell in enumerate(row):
                        row_cells[col_idx].text = cell
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            else:
                for val in item.values:
                    doc.add_paragraph(f"- {val}")
            for att in item.attachments:
                desc = att.content_type or "binary"
                doc.add_paragraph(f"Attachment: {att.name} ({desc})")
            doc.add_paragraph()

        destination.parent.mkdir(parents=True, exist_ok=True)
        doc.save(destination)

    def save_attachments(self, attachments: List[Attachment]) -> List[Path]:
        """Persist decoded attachments to disk."""
        saved = []
        self.output_dir.mkdir(parents=True, exist_ok=True)
        for attachment in attachments:
            target = self.output_dir / attachment.name
            target.write_bytes(attachment.content)
            saved.append(target)
        return saved

    def create_email_message(self, recipient: str, sender: str, attachments: List[Path], poll_id: str, creator: str = None, created: str = None, survey_attachments: List[Path] = None) -> EmailMessage:
        """Build email message with body and attachments."""
        survey_attachments = survey_attachments or []

        # Format email body with creator and created information
        body_lines = ["Survey results generated by SurveyConverter.", ""]
        if creator:
            body_lines.append(f"Created by: {creator}")
        if created:
            # Parse ISO timestamp and format it in a human-readable way
            from datetime import datetime
            try:
                dt = datetime.fromisoformat(created.replace('Z', '+00:00'))
                formatted_date = dt.strftime("%B %d, %Y at %I:%M %p %Z")
                body_lines.append(f"Created on: {formatted_date}")
            except (ValueError, AttributeError):
                # If parsing fails, just use the raw timestamp
                body_lines.append(f"Created on: {created}")

        # Update attachment description in body
        body_lines.append("")
        if survey_attachments:
            body_lines.append(f"Attachments: {len(attachments)} format file(s) and {len(survey_attachments)} survey attachment(s).")
        else:
            body_lines.append("Attachments: all requested formats.")

        message = EmailMessage()
        message["From"] = sender
        message["To"] = recipient
        message["Subject"] = f"SurveyJS export ({poll_id})"
        message["Date"] = email.utils.formatdate(localtime=True)
        message["Message-ID"] = email.utils.make_msgid(domain=sender.split("@")[-1])
        message.set_content("\n".join(body_lines))

        # Attach format files
        for path in attachments:
            data = path.read_bytes()
            ctype, _ = mimetypes.guess_type(path.name)
            maintype, subtype = (ctype or "application/octet-stream").split("/", 1)
            message.add_attachment(data, maintype=maintype, subtype=subtype, filename=path.name)

        # Attach survey attachments (images and binary files)
        for path in survey_attachments:
            data = path.read_bytes()
            ctype, _ = mimetypes.guess_type(path.name)
            maintype, subtype = (ctype or "application/octet-stream").split("/", 1)
            message.add_attachment(data, maintype=maintype, subtype=subtype, filename=path.name)

        return message

    def send_email_smtp(self, message: EmailMessage, recipient: str, host: str, port: int, username: str = None, password: str = None, use_starttls: bool = False) -> None:
        """Send email message via SMTP."""
        # Log SMTP configuration (masking password)
        logger.info("SMTP Configuration:")
        logger.info("  Host: %s", host)
        logger.info("  Port: %s", port)
        logger.info("  Username: %s", username or "(not set)")
        logger.info("  Password: %s", "***" if password else "(not set)")
        logger.info("  Use STARTTLS: %s", use_starttls)

        try:
            logger.info("Sending email to %s via %s:%s", recipient, host, port)
            with smtplib.SMTP(host, port) as smtp:
                if use_starttls:
                    smtp.starttls()
                if username and password:
                    smtp.login(username, password)
                smtp.send_message(message)
            logger.info("Email sent to %s", recipient)
        except Exception:
            logger.exception("Failed to send email to %s via %s:%s", recipient, host, port)
            raise

    def send_email(self, recipient: str, attachments: List[Path], poll_id: str, creator: str = None, created: str = None, survey_attachments: List[Path] = None) -> None:
        """Send generated files and survey attachments via SMTP."""
        load_dotenv()
        if not attachments:
            logger.info("No attachments to send; skipping email to %s", recipient)
            return

        survey_attachments = survey_attachments or []

        # Load SMTP configuration
        host = os.environ.get(ENV_SMTP_HOST, "localhost")
        port = int(os.environ.get(ENV_SMTP_PORT, "25"))
        username = os.environ.get(ENV_SMTP_USERNAME)
        password = os.environ.get(ENV_SMTP_PASSWORD)
        use_starttls = os.environ.get(ENV_SMTP_STARTTLS, "false").lower() in {
            "1",
            "true",
            "yes",
            "on",
        }
        sender = os.environ.get(ENV_EMAIL_SENDER, f"surveyjs@{os.uname().nodename}")

        # Create email message
        message = self.create_email_message(
            recipient, sender, attachments, poll_id, creator, created, survey_attachments
        )

        # Send via SMTP
        self.send_email_smtp(message, recipient, host, port, username, password, use_starttls)

    def run(self, formats: set[str], email_recipient: str | None = None) -> List[Path]:
        """Convert the first survey entry to the requested formats."""
        load_dotenv()
        raw_entry = self.load_first_entry()
        entry = raw_entry.get("result", raw_entry)
        poll_id = self.extract_poll_id(raw_entry) or self.extract_poll_id(entry)

        # Extract creator and created information from raw_entry
        creator = raw_entry.get("user")
        created = raw_entry.get("created")

        items, attachments = self.collect_items(entry, poll_id)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        saved_attachments = self.save_attachments(attachments)

        written_paths: List[Path] = []

        if "text" in formats:
            text_lines = self.build_text(items)
            txt_path = self.output_dir / f"{poll_id}.txt"
            txt_path.write_text("\n".join(text_lines), encoding="utf-8")
            written_paths.append(txt_path)

        need_markdown = bool({"md", "html", "pdf"} & formats)
        markdown_body = None
        if need_markdown:
            markdown_body = self.build_markdown(items, poll_id)

        if "md" in formats and markdown_body is not None:
            md_path = self.output_dir / f"{poll_id}.md"
            md_path.write_text(markdown_body, encoding="utf-8")
            written_paths.append(md_path)

        html_body = None
        need_html = bool({"html", "pdf"} & formats)
        if need_html and markdown_body is not None:
            html_body = markdown(markdown_body, extras=["tables"])
            html_body = self.inline_html_images(html_body, attachments)

        if "html" in formats and html_body is not None:
            html_path = self.output_dir / f"{poll_id}.html"
            html_output = self.wrap_html_output(html_body)
            html_path.write_text(html_output, encoding="utf-8")
            written_paths.append(html_path)

        if "pdf" in formats and html_body is not None:
            pdf_path = self.output_dir / f"{poll_id}.pdf"
            pdf_html = self.wrap_pdf_html(html_body, creator, created)
            self.save_pdf(pdf_html, pdf_path)
            written_paths.append(pdf_path)

        need_tabular = bool({"csv", "xlsx", "docx"} & formats)
        if need_tabular:
            table_rows = self.build_table_rows(items)
            if "csv" in formats:
                csv_path = self.output_dir / f"{poll_id}.csv"
                self.write_csv(table_rows, csv_path)
                written_paths.append(csv_path)
            if "xlsx" in formats:
                xlsx_path = self.output_dir / f"{poll_id}.xlsx"
                self.write_excel(table_rows, xlsx_path)
                written_paths.append(xlsx_path)
            if "docx" in formats:
                docx_path = self.output_dir / f"{poll_id}.docx"
                self.write_docx(items, docx_path, poll_id, creator, created)
                written_paths.append(docx_path)

        if "xml" in formats:
            xml_body = self.build_xml(items, poll_id)
            xml_path = self.output_dir / f"{poll_id}.xml"
            xml_path.write_text(xml_body, encoding="utf-8")
            written_paths.append(xml_path)

        if "json" in formats:
            json_body = self.build_json(items, poll_id, creator, created)
            json_path = self.output_dir / f"{poll_id}.json"
            json_path.write_text(json_body, encoding="utf-8")
            written_paths.append(json_path)

        print(f"Poll ID: {poll_id}")
        if written_paths:
            print("Wrote:")
            for path in written_paths:
                print(f"- {path}")
        else:
            print("No formats selected; nothing written.")
        if saved_attachments:
            print("Attachments:")
            for path in saved_attachments:
                print(f"- {path.name}")

        if email_recipient:
            total_email_attachments = len(written_paths) + len(saved_attachments)
            print(f"Sending email to {email_recipient} with {total_email_attachments} attachment(s) ({len(written_paths)} format files, {len(saved_attachments)} survey files)...")
            self.send_email(email_recipient, written_paths, poll_id, creator, created, saved_attachments)
            print(f"Email sent to {email_recipient}")

        return written_paths


def main() -> None:
    """Entry point: parse args and run the converter."""
    args = parse_args()
    try:
        formats = parse_formats(args.formats)
    except ValueError as exc:
        raise SystemExit(str(exc)) from exc

    data_path = Path(args.data)
    form_path = Path(args.form)
    output_dir = Path(args.output)

    if args.email and not logging.getLogger().handlers:
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s %(levelname)s %(name)s: %(message)s",
        )

    converter = SurveyConverter(data_path, form_path, output_dir)
    converter.run(formats, email_recipient=args.email)


if __name__ == "__main__":
    main()
