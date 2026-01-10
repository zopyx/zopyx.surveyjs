from __future__ import annotations

import json
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from zopyx.surveyjs.converters import (
    common,
    csv_export,
    docx_export,
    html,
    json_export,
    markdown,
    pdf,
    text,
    xlsx_export,
    xml_export,
)
from zopyx.surveyjs.converters.types import Attachment, Item


@pytest.fixture
def image_attachment() -> Attachment:
    return Attachment("photo.png", b"\x89PNG", "image/png", field_label="Photo")


@pytest.fixture
def binary_attachment() -> Attachment:
    return Attachment("document.bin", b"binary data", None, field_label="Binary")


@pytest.fixture
def sample_items(image_attachment: Attachment, binary_attachment: Attachment) -> list[Item]:
    return [
        Item(
            key="text",
            label="Text Question",
            values=["alpha", "beta"],
            attachments=[binary_attachment],
        ),
        Item(
            key="table",
            label="Table Question",
            values=[],
            attachments=[image_attachment],
            table=[["Col A", "Col B"], ["Row 1", "Row 2"]],
            table_columns=[("a", "Column A"), ("b", "Column B")],
        ),
        Item(
            key="matrix",
            label="Matrix",
            values=["ignored"],
            attachments=[],
            field_type="matrixdynamic",
            raw_value=[{"col1": "v1"}],
        ),
    ]


def test_attachment_data_url_and_is_image_guessing() -> None:
    bin_attachment = Attachment("file.bin", b"abc", None)
    url = bin_attachment.data_url()
    assert url.startswith("data:application/octet-stream;base64,YWJj")
    assert not bin_attachment.is_image

    png_attachment = Attachment("img.png", b"\x89PNG", "image/png")
    assert png_attachment.is_image
    assert "image/png;base64" in png_attachment.data_url()


def test_render_text_table_empty_and_padding() -> None:
    assert common.render_text_table([]) == ["(empty)"]

    lines = common.render_text_table([["H1"], ["V1", "V2"]])
    assert lines[0].startswith("H1")
    assert lines[1].startswith("-")
    assert "+" in lines[1]
    assert "V1" in lines[2] and "V2" in lines[2]


def test_render_markdown_table_empty_and_padding() -> None:
    assert common.render_markdown_table([]) == ["(empty)"]

    lines = common.render_markdown_table([["H1"], ["V1", "V2"]])
    assert lines[0] == "| H1 |  |"
    assert lines[1] == "| --- | --- |"
    assert lines[2] == "| V1 | V2 |"


def test_build_table_rows_joins_values_and_attachments() -> None:
    items = [
        Item(
            key="k1",
            label="Label 1",
            values=["v1", "v2"],
            attachments=[Attachment("file.txt", b"", "text/plain")],
        ),
        Item(
            key="k2",
            label="Label 2",
            values=["only"],
            attachments=[Attachment("blob.bin", b"", None)],
        ),
    ]
    rows = common.build_table_rows(items)
    assert rows[0] == ("k1", "Label 1", "v1; v2", "file.txt (text/plain)")
    assert rows[1] == ("k2", "Label 2", "only", "blob.bin (binary)")


def test_inline_html_images_replaces_sources(image_attachment: Attachment) -> None:
    html_body = "<p><img src='photo.png'><img src=\"photo.png\"><img src='other.bin'></p>"
    updated = common.inline_html_images(html_body, [image_attachment])
    assert "data:image/png;base64" in updated
    assert "photo.png" not in updated
    assert "other.bin" in updated


def test_wrap_pdf_html_metadata_insertion_with_heading() -> None:
    html_body = "<h1>Title</h1><p>Body</p>"
    wrapped = common.wrap_pdf_html(
        html_body, creator="Alice", created="2024-05-15T10:20:00Z"
    )
    assert wrapped.startswith("<html>")
    assert "Created by:</strong> Alice" in wrapped
    assert "May 15, 2024 at" in wrapped
    assert wrapped.index("</h1>") < wrapped.index("Created by:")


def test_wrap_pdf_html_without_heading_uses_raw_date() -> None:
    wrapped = common.wrap_pdf_html("<p>Body</p>", created="not-a-date")
    assert wrapped.startswith("<html>")
    assert "not-a-date" in wrapped


def test_wrap_html_output_adds_style() -> None:
    wrapped = common.wrap_html_output("<p>Body</p>")
    assert wrapped.startswith("<html>")
    assert "<style>" in wrapped
    assert "Body" in wrapped


def test_build_text_handles_tables_and_metadata(sample_items: list[Item]) -> None:
    lines = text.build_text(
        sample_items, creator="User", created="2024-05-15T10:20:00Z"
    )
    assert lines[0] == "Survey response"
    assert any("Created by: User" in line for line in lines)
    assert any("Text Question" in line for line in lines)
    assert any("Attachment" in line for line in lines)
    assert any("Col A" in line for line in lines)


def test_build_text_handles_unparseable_date(sample_items: list[Item]) -> None:
    lines = text.build_text(sample_items, created="yesterday")
    assert any("Created on: yesterday" in line for line in lines)


def test_write_text_creates_file(tmp_path: Path, sample_items: list[Item]) -> None:
    dest = tmp_path / "out" / "survey.txt"
    path = text.write_text(sample_items, dest, creator="User")
    assert path == dest
    content = dest.read_text(encoding="utf-8")
    assert "Survey response" in content
    assert dest.exists()


def test_build_markdown_includes_table_and_attachments(
    sample_items: list[Item],
) -> None:
    md = markdown.build_markdown(
        sample_items, poll_id="poll-1", creator="Bob", created="2024-05-15T10:20:00Z"
    )
    assert md.startswith("# Survey response (poll-1)")
    assert "Created by: Bob" in md
    assert "| Col A | Col B |" in md
    assert "![Table Question - photo.png](photo.png)" in md
    assert "[document.bin](document.bin)" in md


def test_write_markdown_creates_file(tmp_path: Path, sample_items: list[Item]) -> None:
    dest = tmp_path / "md" / "survey.md"
    path = markdown.write_markdown(sample_items, "poll-2", dest)
    assert path == dest
    assert dest.read_text(encoding="utf-8").startswith("# Survey response (poll-2)")


def test_build_html_inlines_images_and_tables(image_attachment: Attachment) -> None:
    md_text = "# Title\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n\n![Alt](photo.png)"
    html_body = html.build_html(md_text, [image_attachment])
    assert "<table>" in html_body and "<td>1</td>" in html_body
    assert "data:image/png;base64" in html_body


def test_write_html_wraps_body(tmp_path: Path, image_attachment: Attachment) -> None:
    dest = tmp_path / "html" / "survey.html"
    html.write_html("# Title", [image_attachment], dest)
    content = dest.read_text(encoding="utf-8")
    assert content.startswith("<html>")
    assert "<style>" in content


def test_write_pdf_uses_weasyprint(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    captured = {}

    class DummyHTML:
        def __init__(self, string: str) -> None:
            captured["string"] = string

        def write_pdf(self, destination: Path) -> None:
            destination.write_bytes(b"%PDF-1.4 dummy")
            captured["destination"] = destination

    monkeypatch.setattr(pdf, "HTML", DummyHTML)

    dest = tmp_path / "pdf" / "survey.pdf"
    result = pdf.write_pdf("<h1>Hi</h1>", dest, creator="Alice")
    assert result == dest
    assert dest.exists()
    assert "%PDF" in dest.read_bytes().decode("latin-1")
    assert "Created by" in captured["string"]


def test_write_csv_outputs_rows(tmp_path: Path) -> None:
    dest = tmp_path / "csv" / "survey.csv"
    rows = [("k", "label", "value", "att")]
    path = csv_export.write_csv(rows, dest)
    assert path == dest
    content = dest.read_text(encoding="utf-8").splitlines()
    assert content[0] == "Key,Field,Value,Attachments"
    assert content[1] == "k,label,value,att"


def test_write_xlsx_outputs_rows(tmp_path: Path) -> None:
    dest = tmp_path / "xlsx" / "survey.xlsx"
    rows = [("k", "label", "value", "att")]
    xlsx_export.write_xlsx(rows, dest)
    wb = load_workbook(dest)
    ws = wb.active
    assert ws.title == "Survey"
    assert ws.cell(row=1, column=1).value == "Key"
    assert ws.cell(row=2, column=2).value == "label"


def test_build_xml_handles_tables_and_attachments(
    sample_items: list[Item], image_attachment: Attachment
) -> None:
    items = [
        sample_items[0],
        Item(
            key="table",
            label="Table Question",
            values=[],
            attachments=[image_attachment],
            table=[["H1", "H2", "H3"], ["A", "B", "C"]],
            table_columns=[("col1", "Column 1")],
        ),
    ]
    xml_text = xml_export.build_xml(items, "poll-xml")
    root = ET.fromstring(xml_text.split("\n", 1)[1])
    assert root.tag == "survey_response"
    assert root.get("poll_id") == "poll-xml"
    field = root.findall("field")[1]
    table = field.find("table")
    assert table is not None
    header_row = table.findall("row")[0]
    assert header_row.get("header") == "true"
    cells = header_row.findall("cell")
    assert cells[0].get("label") == "Column 1"
    assert cells[1].get("label") is None
    attachments = field.find("attachments")
    assert attachments is not None
    attachment = attachments.find("attachment")
    assert attachment.get("is_image") == "true"


def test_write_xml_creates_file(tmp_path: Path, sample_items: list[Item]) -> None:
    dest = tmp_path / "xml" / "survey.xml"
    path = xml_export.write_xml(sample_items, "poll-xml", dest)
    assert path == dest
    assert dest.exists()
    assert dest.read_text(encoding="utf-8").startswith('<?xml version="1.0"')


def test_write_docx_writes_content(tmp_path: Path, sample_items: list[Item]) -> None:
    dest = tmp_path / "docx" / "survey.docx"
    docx_export.write_docx(sample_items, dest, poll_id="poll-doc")
    doc = Document(dest)
    assert doc.paragraphs[0].text == "Survey Response: poll-doc"
    texts = [p.text for p in doc.paragraphs if p.text]
    assert any("Text Question" in t for t in texts)
    assert any("Attachment: document.bin" in t for t in texts)
    assert doc.tables[0].cell(0, 0).text == "Col A"


def test_build_json_handles_matrix_values(sample_items: list[Item]) -> None:
    json_text = json_export.build_json(
        sample_items, poll_id="poll-json", creator="Creator", created="2024-05-15"
    )
    payload = json.loads(json_text)
    assert payload["poll_id"] == "poll-json"
    assert payload["creator"] == "Creator"
    matrix_field = payload["fields"][2]
    assert matrix_field["values"] == [{"col1": "v1"}]
    attachment_meta = payload["fields"][1]["attachments"][0]
    assert attachment_meta["is_image"] is True


def test_write_json_creates_file(tmp_path: Path, sample_items: list[Item]) -> None:
    dest = tmp_path / "json" / "survey.json"
    path = json_export.write_json(sample_items, "poll-json", dest)
    assert path == dest
    data = json.loads(dest.read_text(encoding="utf-8"))
    assert data["poll_id"] == "poll-json"
